"""
增强版文档解析器（来自 RAG-Pro）
支持格式：PDF, DOCX, DOC, CSV, TXT, MD
相比基础版的优势：
- 更好的 PDF 解析（PyMuPDF）
- DOCX 标题结构保留
- CSV 结构化解析
- 多种编码自动识别
- Markdown 标题层级识别
"""

import csv
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple


# 尝试导入可选依赖
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("Warning: PyMuPDF not installed, PDF parsing will use fallback")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed, DOCX parsing will use fallback")

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False
    print("Warning: chardet not installed, encoding detection limited")


@dataclass
class DocumentPage:
    """文档页面结构"""
    text: str
    page_number: int | None = None
    section_title: str | None = None
    metadata: dict = field(default_factory=dict)


class DocumentParserV2:
    """增强版文档解析器"""

    PARSERS = {
        "pdf": "_parse_pdf",
        "docx": "_parse_docx",
        "doc": "_parse_docx",
        "csv": "_parse_csv",
        "txt": "_parse_txt",
        "md": "_parse_md",
    }

    def __init__(self):
        pass

    def parse_file(self, file_path: str) -> List[DocumentPage]:
        """
        从文件路径解析文档
        
        Args:
            file_path: 文件路径
        
        Returns:
            DocumentPage 列表
        """
        path = Path(file_path)
        ext = path.suffix.lstrip(".").lower()
        parser_method = self.PARSERS.get(ext)
        
        if not parser_method:
            raise ValueError(f"Unsupported file type: {ext}")
        
        return getattr(self, parser_method)(file_path)

    def parse_bytes(self, content: bytes, filename: str) -> List[DocumentPage]:
        """
        从字节内容解析文档
        
        Args:
            content: 文件字节内容
            filename: 文件名（用于判断格式）
        
        Returns:
            DocumentPage 列表
        """
        ext = Path(filename).suffix.lstrip(".").lower()
        
        if ext == "pdf":
            return self._parse_pdf_bytes(content)
        elif ext in ["docx", "doc"]:
            return self._parse_docx_bytes(content)
        elif ext == "csv":
            return self._parse_csv_bytes(content, filename)
        elif ext in ["txt", "md"]:
            return self._parse_text_bytes(content, filename)
        else:
            return self._parse_text_bytes(content, filename)

    def _parse_pdf(self, file_path: str) -> List[DocumentPage]:
        """解析 PDF（使用 PyMuPDF）"""
        pages = []
        
        if HAS_PYMUPDF:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text")
                if text.strip():
                    pages.append(DocumentPage(
                        text=text.strip(),
                        page_number=page_num,
                        metadata={"source": Path(file_path).name},
                    ))
            doc.close()
        else:
            pages = self._parse_pdf_fallback(file_path)
        
        return pages

    def _parse_pdf_bytes(self, content: bytes) -> List[DocumentPage]:
        """从字节解析 PDF"""
        pages = []
        
        if HAS_PYMUPDF:
            doc = fitz.open(stream=content, filetype="pdf")
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text")
                if text.strip():
                    pages.append(DocumentPage(
                        text=text.strip(),
                        page_number=page_num,
                        metadata={"source": "bytes"},
                    ))
            doc.close()
        else:
            pages = self._parse_pdf_bytes_fallback(content)
        
        return pages

    def _parse_pdf_fallback(self, file_path: str) -> List[DocumentPage]:
        """PDF 解析降级方案（PyPDF2）"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text:
                    pages.append(DocumentPage(
                        text=text,
                        page_number=page_num,
                        metadata={"source": Path(file_path).name},
                    ))
            return pages
        except Exception as e:
            return [DocumentPage(
                text=f"[PDF 解析失败: {str(e)}]",
                page_number=1,
                metadata={"source": Path(file_path).name, "error": str(e)},
            )]

    def _parse_pdf_bytes_fallback(self, content: bytes) -> List[DocumentPage]:
        """PDF 字节解析降级方案"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages = []
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text:
                    pages.append(DocumentPage(
                        text=text,
                        page_number=page_num,
                        metadata={"source": "bytes"},
                    ))
            return pages
        except Exception as e:
            return [DocumentPage(
                text=f"[PDF 解析失败: {str(e)}]",
                page_number=1,
                metadata={"source": "bytes", "error": str(e)},
            )]

    def _parse_docx(self, file_path: str) -> List[DocumentPage]:
        """解析 DOCX，保留标题结构"""
        pages = []
        
        if HAS_DOCX:
            doc = DocxDocument(file_path)
            current_section = None
            current_text_parts: List[str] = []
            page_counter = 1

            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    if current_text_parts:
                        pages.append(DocumentPage(
                            text="\n".join(current_text_parts).strip(),
                            page_number=page_counter,
                            section_title=current_section,
                            metadata={"source": Path(file_path).name},
                        ))
                        page_counter += 1
                        current_text_parts = []
                    current_section = para.text.strip()
                    current_text_parts.append(para.text)
                elif para.text.strip():
                    current_text_parts.append(para.text)

            if current_text_parts:
                pages.append(DocumentPage(
                    text="\n".join(current_text_parts).strip(),
                    page_number=page_counter,
                    section_title=current_section,
                    metadata={"source": Path(file_path).name},
                ))

            if not pages:
                full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if full_text:
                    pages.append(DocumentPage(
                        text=full_text,
                        page_number=1,
                        metadata={"source": Path(file_path).name},
                    ))
        else:
            pages = self._parse_docx_fallback(file_path)
        
        return pages

    def _parse_docx_bytes(self, content: bytes) -> List[DocumentPage]:
        """从字节解析 DOCX"""
        pages = []
        
        if HAS_DOCX:
            doc = DocxDocument(io.BytesIO(content))
            current_section = None
            current_text_parts: List[str] = []
            page_counter = 1

            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    if current_text_parts:
                        pages.append(DocumentPage(
                            text="\n".join(current_text_parts).strip(),
                            page_number=page_counter,
                            section_title=current_section,
                            metadata={"source": "bytes"},
                        ))
                        page_counter += 1
                        current_text_parts = []
                    current_section = para.text.strip()
                    current_text_parts.append(para.text)
                elif para.text.strip():
                    current_text_parts.append(para.text)

            if current_text_parts:
                pages.append(DocumentPage(
                    text="\n".join(current_text_parts).strip(),
                    page_number=page_counter,
                    section_title=current_section,
                    metadata={"source": "bytes"},
                ))
        else:
            pages = [DocumentPage(text=str(content), page_number=1, metadata={"source": "bytes"})]
        
        return pages

    def _parse_docx_fallback(self, file_path: str) -> List[DocumentPage]:
        """DOCX 解析降级方案"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return [DocumentPage(
                text=text,
                page_number=1,
                metadata={"source": Path(file_path).name},
            )]
        except Exception:
            return [DocumentPage(text="[DOCX 解析失败]", page_number=1, metadata={"source": Path(file_path).name})]

    def _parse_csv(self, file_path: str) -> List[DocumentPage]:
        """解析 CSV 为结构化文本"""
        pages = []
        content, used_encoding = self._read_text_file(file_path)
        if content is None:
            raise ValueError("Unable to decode CSV file with supported encodings")

        reader = csv.DictReader(io.StringIO(content))
        rows = list(reader)

        if not rows:
            return pages

        headers = list(rows[0].keys())
        batch_size = 20

        for i in range(0, len(rows), batch_size):
            batch = rows[i:i + batch_size]
            text_parts = []
            for row in batch:
                row_text = "; ".join(f"{k}: {v}" for k, v in row.items() if v)
                text_parts.append(row_text)

            pages.append(DocumentPage(
                text="\n".join(text_parts),
                page_number=(i // batch_size) + 1,
                section_title=f"数据行 {i + 1}-{i + len(batch)}",
                metadata={
                    "source": Path(file_path).name,
                    "headers": headers,
                    "total_rows": len(rows),
                    "encoding": used_encoding,
                },
            ))

        return pages

    def _parse_csv_bytes(self, content: bytes, filename: str) -> List[DocumentPage]:
        """从字节解析 CSV"""
        pages = []
        content_str, used_encoding = self._decode_bytes(content)
        
        reader = csv.DictReader(io.StringIO(content_str))
        rows = list(reader)

        if not rows:
            return pages

        headers = list(rows[0].keys())
        batch_size = 20

        for i in range(0, len(rows), batch_size):
            batch = rows[i:i + batch_size]
            text_parts = []
            for row in batch:
                row_text = "; ".join(f"{k}: {v}" for k, v in row.items() if v)
                text_parts.append(row_text)

            pages.append(DocumentPage(
                text="\n".join(text_parts),
                page_number=(i // batch_size) + 1,
                section_title=f"数据行 {i + 1}-{i + len(batch)}",
                metadata={
                    "source": filename,
                    "headers": headers,
                    "total_rows": len(rows),
                    "encoding": used_encoding,
                },
            ))

        return pages

    def _parse_txt(self, file_path: str) -> List[DocumentPage]:
        """解析文本文件"""
        text, used_encoding = self._read_text_file(file_path)
        if text is None:
            raise ValueError("Unable to decode text file with supported encodings")

        if not text.strip():
            return []

        sections = text.split("\n\n")
        pages = []
        current_text_parts: List[str] = []
        page_counter = 1
        char_count = 0

        for section in sections:
            section = section.strip()
            if not section:
                continue
            current_text_parts.append(section)
            char_count += len(section)

            if char_count >= 2000:
                pages.append(DocumentPage(
                    text="\n\n".join(current_text_parts),
                    page_number=page_counter,
                    metadata={
                        "source": Path(file_path).name,
                        "encoding": used_encoding,
                    },
                ))
                current_text_parts = []
                char_count = 0
                page_counter += 1

        if current_text_parts:
            pages.append(DocumentPage(
                text="\n\n".join(current_text_parts),
                page_number=page_counter,
                metadata={
                    "source": Path(file_path).name,
                    "encoding": used_encoding,
                },
            ))

        return pages

    def _parse_md(self, file_path: str) -> List[DocumentPage]:
        """解析 Markdown，保留标题结构"""
        text, used_encoding = self._read_text_file(file_path)
        if text is None:
            raise ValueError("Unable to decode markdown file with supported encodings")

        if not text.strip():
            return []

        lines = text.split("\n")
        pages = []
        current_section = None
        current_text_parts: List[str] = []
        page_counter = 1

        for line in lines:
            if line.startswith("# ") or line.startswith("## "):
                if current_text_parts:
                    pages.append(DocumentPage(
                        text="\n".join(current_text_parts).strip(),
                        page_number=page_counter,
                        section_title=current_section,
                        metadata={
                            "source": Path(file_path).name,
                            "encoding": used_encoding,
                        },
                    ))
                    page_counter += 1
                    current_text_parts = []
                current_section = line.lstrip("# ").strip()
            current_text_parts.append(line)

        if current_text_parts:
            pages.append(DocumentPage(
                text="\n".join(current_text_parts).strip(),
                page_number=page_counter,
                section_title=current_section,
                metadata={
                    "source": Path(file_path).name,
                    "encoding": used_encoding,
                },
            ))

        return pages

    def _parse_text_bytes(self, content: bytes, filename: str) -> List[DocumentPage]:
        """解析文本字节内容"""
        text, used_encoding = self._decode_bytes(content)
        
        if not text.strip():
            return []

        sections = text.split("\n\n")
        pages = []
        current_text_parts: List[str] = []
        page_counter = 1
        char_count = 0

        for section in sections:
            section = section.strip()
            if not section:
                continue
            current_text_parts.append(section)
            char_count += len(section)

            if char_count >= 2000:
                pages.append(DocumentPage(
                    text="\n\n".join(current_text_parts),
                    page_number=page_counter,
                    metadata={
                        "source": filename,
                        "encoding": used_encoding,
                    },
                ))
                current_text_parts = []
                char_count = 0
                page_counter += 1

        if current_text_parts:
            pages.append(DocumentPage(
                text="\n\n".join(current_text_parts),
                page_number=page_counter,
                metadata={
                    "source": filename,
                    "encoding": used_encoding,
                },
            ))

        return pages

    def _read_text_file(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """读取文本文件，尝试多种编码"""
        raw = Path(file_path).read_bytes()
        return self._decode_bytes(raw)

    def _decode_bytes(self, raw: bytes) -> Tuple[str, Optional[str]]:
        """解码字节内容，尝试多种编码"""
        encodings: List[str] = []
        
        if HAS_CHARDET:
            detected = chardet.detect(raw)
            if detected.get("encoding"):
                encodings.append(detected["encoding"])

        encodings.extend([
            "utf-8-sig",
            "utf-8",
            "utf-16",
            "utf-16-le",
            "utf-16-be",
            "gb18030",
            "gbk",
            "gb2312",
            "big5",
            "cp936",
            "latin-1",
        ])

        seen: set[str] = set()
        for encoding in encodings:
            if not encoding:
                continue
            encoding = encoding.lower()
            if encoding in seen:
                continue
            seen.add(encoding)
            try:
                return raw.decode(encoding), encoding
            except UnicodeDecodeError:
                continue

        return raw.decode("utf-8", errors="replace"), "utf-8(replace)"


# 单例实例
document_parser_v2 = DocumentParserV2()
